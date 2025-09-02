#!/usr/bin/env python3
import argparse
import json
import textwrap
from pathlib import Path
from typing import Dict, List, Optional

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages


NARRATIVE = {
    'title': 'Learning Signatures of Good Writing — Final Report',
    'summary': (
        "Great prose and poetry ride a cadence: mostly focused choices, punctuated by purposeful spikes of surprise "
        "that turn the scene or idea, followed by a short cooldown that grounds what just happened. Spikes align with "
        "content words and rhetorical pivots (not punctuation), with larger sustained shifts every few sentences or lines.\n\n"
        "Our analysis measured token-level distributions (p_true, entropy, rank, nucleus width), cadence statistics "
        "(spike rate, inter-peak intervals, cooldown entropy drop), cohesion (order vs shuffled), and token-class contexts. "
        "We then built a cadence-aware sampler (HF + MLX) that enforces per-phase top_p/temperature, content-aware spikes, "
        "cooldowns, and optional rhyme/line nudges."
    ),
    'principles': [
        'Cadence, not chaos: base focus → spike → cooldown → repeat',
        'Spike on content pivots; defer punctuation/newline',
        'Sustained shifts every 1–3 lines/sentences',
        'Order matters: negative cohesion delta (original > shuffled)',
        'Genre dials: denser spikes for poetry; gentler cadence for prose',
    ],
    'howto': (
        "How to read the signatures: Surprisal/entropy show focus vs openness; spike rate and IPI (inter-peak interval) capture rhythm; "
        "cooldown entropy drop shows consolidation after turns; cohesion delta quantifies how much word order matters in an author; "
        "content vs punctuation alignment around spikes shows whether turns land on meaningful tokens."
    ),
    'findings': (
        "Findings by genre and author:\n"
        "- Sonnets: medium spike density, clear quatrain cadence, a volta near line 9; strong end-words and rhyme.\n"
        "- Dickinson: higher micro-turn density with short cooldowns; punctuation tolerance (dashes) around content spikes.\n"
        "- Imagist free verse: moderate spike density but stronger spike intensity; concrete nouns/verbs; sparse function words.\n"
        "- Whitman/long-line: lower spike density with stanza-level sustained shifts; broader swell windows rather than frequent micro-spikes.\n"
        "- Wodehouse: playful but coherent cadence; frequent callbacks and entity threading; negative cohesion delta is notable.\n"
        "- Hemingway: simpler clauses with steadier cadence; fewer spikes, stronger cooldown consolidation; high cohesion."
    ),
    'best_practices': [
        'Base top_p ≈ 0.88–0.92, temperature ≈ 0.7–0.85',
        'Spike top_p ≈ 0.95–0.98, temperature ≈ 0.95–1.08; require content token',
        'Cooldown top_p ≈ 0.80–0.86, temperature ≈ 0.6–0.75 for 3–8 tokens',
        'Target spike interval: poetry 8–16; prose 14–24 (tune by author)',
        'After spike: look for cooldown entropy drop ≥ 1.0 bits (3-token window)',
        'Maintain content alignment: spike-next-content-rate high; avoid back-to-back punctuation surprises',
    ],
    'what_why': (
        "What we did and why: We measured token distributions and cadence signatures across authors and genres to learn what "
        "‘good writing’ looks like in terms of focus, rhythm, and cohesion. We then built a cadence‑aware sampler that explicitly "
        "follows those patterns — keeping baseline choices focused, inserting purposeful spikes on content pivots, cooling down to consolidate, "
        "and periodically opening sustained shifts. We added rhyme nudging for poetry, repetition controls and a diversity bonus on spikes to keep turns fresh, "
        "and we save before/after snippets so you can see and measure the effect."
    ),
}


def _find_signatures_confusion(model: str) -> Optional[Path]:
    """Locate the most recent signatures confusion matrix PNG for the given model.
    Looks under reports/signatures_* and matches 'gpt2' or 'qwen' based on model name.
    """
    root = Path('reports')
    want = 'gpt2' if 'gpt2' in model.lower() else ('qwen' if 'qwen' in model.lower() else None)
    cands: List[Path] = []
    for d in root.glob('signatures_*'):
        if not d.is_dir():
            continue
        name = d.name.lower()
        if want and want not in name:
            continue
        rep = d / 'classification_report.json'
        fig = d / 'classification_confusion_matrix.png'
        if rep.exists() and fig.exists():
            cands.append(fig)
    if not cands:
        return None
    # Sort by report mtime
    cands.sort(key=lambda p: (p.parent / 'classification_report.json').stat().st_mtime, reverse=True)
    return cands[0]


def load_generated(index_path: Path, model_sub: str) -> List[Dict]:
    rows: List[Dict] = []
    if not index_path.exists():
        return rows
    with index_path.open('r', encoding='utf-8') as f:
        for line in f:
            try:
                e = json.loads(line)
            except Exception:
                continue
            if model_sub.lower() in str(e.get('model','')).lower():
                rows.append(e)
    return rows


def pick_latest_by_preset(entries: List[Dict], presets: List[str], max_per: int = 3) -> Dict[str, List[Dict]]:
    buckets: Dict[str, List[Dict]] = {p: [] for p in presets}
    for e in entries:
        p = e.get('preset')
        if p in buckets:
            buckets[p].append(e)
    # keep latest max_per per preset
    for p in list(buckets.keys()):
        buckets[p] = buckets[p][-max_per:]
    return buckets


def render_final_readme(model: str, model_dir: Path, compare_dir: Path, gen_latest: Dict[str, List[Dict]], out_path: Path):
    lines: List[str] = []
    lines.append(f"# {NARRATIVE['title']}")
    lines.append('')
    lines.append('## Executive Summary')
    lines.append('')
    lines.append(NARRATIVE['summary'])
    lines.append('')
    lines.append('## Principles')
    for b in NARRATIVE['principles']:
        lines.append(f"- {b}")
    lines.append('')
    lines.append('## How to Read These Signatures')
    lines.append('')
    lines.append(NARRATIVE['howto'])
    lines.append('')
    lines.append('## Findings by Genre and Author')
    lines.append('')
    lines.append(NARRATIVE['findings'])
    lines.append('')
    lines.append('## Best Practices — Target Bands')
    for b in NARRATIVE['best_practices']:
        lines.append(f"- {b}")
    lines.append('')

    # Illustrations
    lines.append('## Illustrations')
    import os
    # Always include key compare figures
    illos: List[Path] = []
    compare_picks = [
        'authors_delta_surprisal_mean_mean.png',
        'authors_delta_cooldown_entropy_drop_3_mean.png',
        'authors_delta_nucleus_w_mean_mean.png',
        'authors_delta_spike_next_content_rate_mean.png',
    ]
    for name in compare_picks:
        p = compare_dir / name
        if p.exists():
            illos.append(p)
    # Add model-level overview and author docs
    model_picks = [
        'authors_entropy_vs_surprisal.png',
        'docs_entropy_vs_surprisal.png',
        'docs_nucleus_width_by_type.png',
        'docs_top_cohesion_delta.png',
        'author_william_shakespeare_series.png',
        'author_william_shakespeare_ipi_hist.png',
        'author_william_shakespeare_surprisal_hist.png',
        'author_p_g_wodehouse_series.png',
        'author_p_g_wodehouse_ipi_hist.png',
        'author_p_g_wodehouse_surprisal_hist.png',
        'doc_ts_poem_william_shakespeare_sonnet_130.png',
        'doc_ts_novel_ernest_hemingway_the_sun_also_rises.png',
        'doc_ts_shortstory_aesop_hercules_and_the_wagoner.png',
    ]
    for name in model_picks:
        p = model_dir / name
        if p.exists():
            illos.append(p)

    # Include signatures confusion matrix if available
    sig_fig = _find_signatures_confusion(model)
    if sig_fig:
        illos.insert(0, sig_fig)

    # Render images as HTML with full width to avoid tiny corner rendering
    for img in illos:
        try:
            rel = os.path.relpath(img.as_posix(), start=out_path.parent.as_posix())
        except Exception:
            rel = img.as_posix()
        lines.append(f'<img src="{rel}" alt="{img.name}" style="max-width: 100%; width: 100%; height: auto;" />')
        lines.append('')

    # What we did and why
    lines.append('## What We Did and Why')
    lines.append('')
    lines.append(NARRATIVE['what_why'])
    lines.append('')

    # Generated snippets
    lines.append('## Generated Snippets (Normal vs Fixed-Up)')
    for preset in ['imagist','sonnet','couplets','prose']:
        entries = gen_latest.get(preset, [])
        if not entries:
            continue
        lines.append(f"### {preset.title()}")
        for idx, e in enumerate(entries, 1):
            prompt = e.get('prompt','').rstrip()
            base_path = e.get('paths',{}).get('baseline')
            fix_path = e.get('paths',{}).get('fixed')
            base_text = Path(base_path).read_text(encoding='utf-8').strip() if base_path and Path(base_path).exists() else ''
            fix_text = Path(fix_path).read_text(encoding='utf-8').strip() if fix_path and Path(fix_path).exists() else ''
            lines.append('')
            lines.append(f"#### Sample {idx}")
            lines.append('')
            lines.append('> ' + prompt.replace('\n','\n> '))
            lines.append('')
            if base_text:
                lines.append('Normal')
                lines.append('')
                lines.append('```')
                lines.append(base_text)
                lines.append('```')
                lines.append('')
            if fix_text:
                lines.append('Fixed-Up (Cadence-Controlled)')
                lines.append('')
                lines.append('```')
                lines.append(fix_text)
                lines.append('```')
                lines.append('')

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text('\n'.join(lines) + '\n', encoding='utf-8')
    print(f"Wrote {out_path}")


def build_final_pdf(readme_path: Path, images: List[Path], gen_latest: Dict[str, List[Dict]], out_pdf: Path):
    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    with PdfPages(out_pdf) as pdf:
        # Title page
        fig = plt.figure(figsize=(8.5, 11))
        fig.text(0.5, 0.7, NARRATIVE['title'], ha='center', va='center', fontsize=20)
        fig.text(0.5, 0.62, 'Final narrative, illustrations, and generated snippets', ha='center', fontsize=11)
        pdf.savefig(fig); plt.close(fig)

        # Narrative page(s)
        def add_text_page(title: str, text: str):
            fig = plt.figure(figsize=(8.5, 11))
            fig.text(0.5, 0.95, title, ha='center', fontsize=14)
            wrapped = textwrap.fill(text, width=90)
            fig.text(0.07, 0.9, wrapped, va='top', fontsize=10)
            pdf.savefig(fig); plt.close(fig)

        add_text_page('Executive Summary', NARRATIVE['summary'])
        add_text_page('Principles', '\n'.join('- ' + p for p in NARRATIVE['principles']))

        # Illustration pages
        for img in images:
            if not img.exists():
                continue
            fig = plt.figure(figsize=(8.5, 11))
            ax = fig.add_axes([0.05, 0.12, 0.9, 0.8])
            try:
                im = plt.imread(img)
            except Exception:
                continue
            ax.imshow(im, aspect='auto')
            ax.set_xticks([]); ax.set_yticks([])
            ax.set_frame_on(False)
            cap = img.name.replace('_', ' ').replace('.png', '')
            fig.text(0.5, 0.05, cap, ha='center', fontsize=10)
            pdf.savefig(fig); plt.close(fig)

        # Generated snippets: 1–2 pages per preset with fixed-up text (truncate if needed)
        for preset in ['imagist','sonnet','couplets','prose']:
            entries = gen_latest.get(preset, [])
            for e in entries[-2:]:  # last 2 per preset
                fix_path = e.get('paths',{}).get('fixed')
                if not fix_path or not Path(fix_path).exists():
                    continue
                txt = Path(fix_path).read_text(encoding='utf-8').strip()
                if len(txt) > 2200:
                    txt = txt[:2200] + '...'
                fig = plt.figure(figsize=(8.5, 11))
                fig.text(0.5, 0.95, f"Fixed-Up — {preset.title()}", ha='center', fontsize=14)
                wrapped = textwrap.fill(txt, width=90)
                fig.text(0.07, 0.9, wrapped, va='top', fontsize=10)
                pdf.savefig(fig); plt.close(fig)
    print(f"PDF written to {out_pdf}")


def build_final_html(model: str, images: List[Path], gen_latest: Dict[str, List[Dict]], out_html: Path):
    import base64
    out_html.parent.mkdir(parents=True, exist_ok=True)

    def b64_img(path: Path) -> Optional[str]:
        try:
            data = path.read_bytes()
            enc = base64.b64encode(data).decode('ascii')
            mime = 'image/png'
            return f"data:{mime};base64,{enc}"
        except Exception:
            return None

    html = []
    html.append("<!doctype html>")
    html.append("<html lang=\"en\">\n<head>\n<meta charset=\"utf-8\">\n<title>Final Report</title>")
    html.append("<style>body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;margin:24px;line-height:1.5;} h1,h2,h3{margin-top:1.4em;} .img{margin:18px 0;} img{max-width:100%;height:auto;display:block;} .prompt{background:#f6f8fa;padding:8px 10px;border-left:4px solid #0366d6;white-space:pre-wrap;} pre{background:#0f1117;color:#e6edf3;padding:12px;border-radius:6px;overflow:auto;} .two-col{display:grid;grid-template-columns:1fr 1fr;gap:12px;} .caption{font-size:12px;color:#666;margin-top:4px;} .sample{margin-bottom:24px;border-bottom:1px solid #eee;padding-bottom:16px;}</style>")
    html.append("</head><body>")

    # Title
    html.append(f"<h1>{NARRATIVE['title']}</h1>")
    html.append("<h2>Executive Summary</h2>")
    html.append(f"<p>{NARRATIVE['summary']}</p>")
    html.append("<h2>Principles</h2><ul>")
    for p in NARRATIVE['principles']:
        html.append(f"<li>{p}</li>")
    html.append("</ul>")
    html.append("<h2>How to Read These Signatures</h2>")
    html.append(f"<p>{NARRATIVE['howto']}</p>")
    html.append("<h2>Findings by Genre and Author</h2>")
    html.append('<div style="white-space:pre-wrap">' + NARRATIVE['findings'] + '</div>')
    html.append("<h2>Best Practices — Target Bands</h2><ul>")
    for p in NARRATIVE['best_practices']:
        html.append(f"<li>{p}</li>")
    html.append("</ul>")

    # Illustrations
    html.append("<h2>Illustrations</h2>")
    for img in images:
        if not img.exists():
            continue
        data = b64_img(img)
        if data:
            html.append('<div class="img">')
            html.append(f"<img src=\"{data}\" alt=\"{img.name}\">")
            html.append(f"<div class=\"caption\">{img.name.replace('_',' ').replace('.png','')}</div>")
            html.append("</div>")

    # Snippets
    html.append("<h2>Generated Snippets (Normal vs Fixed-Up)</h2>")
    for preset in ['imagist','sonnet','couplets','prose']:
        entries = gen_latest.get(preset, [])
        if not entries:
            continue
        html.append(f"<h3>{preset.title()}</h3>")
        for idx, e in enumerate(entries, 1):
            prompt = e.get('prompt','').rstrip()
            base_path = e.get('paths',{}).get('baseline')
            fix_path = e.get('paths',{}).get('fixed')
            base_text = Path(base_path).read_text(encoding='utf-8').strip() if base_path and Path(base_path).exists() else ''
            fix_text = Path(fix_path).read_text(encoding='utf-8').strip() if fix_path and Path(fix_path).exists() else ''
            html.append('<div class="sample">')
            html.append(f"<h4>Sample {idx}</h4>")
            html.append('<div class="prompt">' + prompt.replace('<','&lt;').replace('>','&gt;') + '</div>')
            html.append('<div class="two-col">')
            # Normal
            html.append('<div>')
            html.append('<div class="caption">Normal</div>')
            html.append('<pre>' + base_text.replace('<','&lt;').replace('>','&gt;') + '</pre>')
            html.append('</div>')
            # Fixed
            html.append('<div>')
            html.append('<div class="caption">Fixed-Up (Cadence-Controlled)</div>')
            html.append('<pre>' + fix_text.replace('<','&lt;').replace('>','&gt;') + '</pre>')
            html.append('</div>')
            html.append('</div>')
            html.append('</div>')

    html.append("</body></html>")
    out_html.write_text("\n".join(html), encoding='utf-8')
    print(f"HTML written to {out_html}")


def main():
    ap = argparse.ArgumentParser(description='Create final README and PDF combining illustrations and generated snippets')
    ap.add_argument('--model', default='Qwen/Qwen2.5-1.5B')
    ap.add_argument('--out-readme', default='reports/final/README.md')
    ap.add_argument('--out-pdf', default='reports/final/report.pdf')
    ap.add_argument('--out-docx', default='reports/final/report.docx')
    ap.add_argument('--keep-pdf', action='store_true', help='Keep the generated PDF (by default it is deleted)')
    args = ap.parse_args()

    # Resolve model report dir and compare dir
    model_safe = args.model.replace('/', '/')
    model_leaf = args.model.split('/')[-1]
    model_dir = Path('reports') / (args.model.split('/')[0]) / model_leaf
    compare_dir = Path('reports') / f"compare_gpt2_vs_{args.model.replace('/', '_')}"

    # Load generated latest entries per preset for this model
    gen_entries = load_generated(Path('data/generated/index.jsonl'), args.model)
    latest = pick_latest_by_preset(gen_entries, ['imagist','sonnet','couplets','prose'], max_per=3)

    # Build README
    render_final_readme(args.model, model_dir, compare_dir, latest, Path(args.out_readme))

    # Build PDF
    images = [
        compare_dir / 'authors_delta_surprisal_mean_mean.png',
        compare_dir / 'authors_delta_cooldown_entropy_drop_3_mean.png',
        compare_dir / 'authors_delta_nucleus_w_mean_mean.png',
        compare_dir / 'authors_delta_spike_next_content_rate_mean.png',
        model_dir / 'authors_entropy_vs_surprisal.png',
        model_dir / 'docs_nucleus_width_by_type.png',
        model_dir / 'author_william_shakespeare_series.png',
        model_dir / 'author_p_g_wodehouse_series.png',
    ]
    out_pdf_path = Path(args.out_pdf)
    build_final_pdf(Path(args.out_readme), images, latest, out_pdf_path)
    # Delete PDF unless requested to keep
    if not args.keep_pdf:
        try:
            out_pdf_path.unlink(missing_ok=True)
            print(f"Deleted PDF (as requested): {out_pdf_path}")
        except Exception as e:
            print(f"[WARN] Could not delete PDF {out_pdf_path}: {e}")

    # Build DOCX (best-effort)
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore
        docx_path = Path(args.out_docx)
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading(NARRATIVE['title'], 0)
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(NARRATIVE['summary'])
        doc.add_heading('Principles', level=1)
        for p in NARRATIVE['principles']:
            doc.add_paragraph(p, style='List Bullet')
        doc.add_heading('How to Read These Signatures', level=1)
        doc.add_paragraph(NARRATIVE['howto'])
        doc.add_heading('Findings by Genre and Author', level=1)
        for line in NARRATIVE['findings'].split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_heading('Best Practices — Target Bands', level=1)
        for p in NARRATIVE['best_practices']:
            doc.add_paragraph(p, style='List Bullet')
        # Images
        doc.add_heading('Illustrations', level=1)
        for img in images:
            if img.exists():
                doc.add_paragraph(img.name)
                try:
                    doc.add_picture(str(img), width=Inches(6.5))
                except Exception:
                    pass
        # Snippets
        doc.add_heading('Generated Snippets (Normal vs Fixed-Up)', level=1)
        for preset in ['imagist','sonnet','couplets','prose']:
            entries = latest.get(preset, [])
            if not entries:
                continue
            doc.add_heading(preset.title(), level=2)
            for idx, e in enumerate(entries, 1):
                doc.add_heading(f'Sample {idx}', level=3)
                prompt = e.get('prompt','').strip()
                base_path = e.get('paths',{}).get('baseline')
                fix_path = e.get('paths',{}).get('fixed')
                base_text = Path(base_path).read_text(encoding='utf-8').strip() if base_path and Path(base_path).exists() else ''
                fix_text = Path(fix_path).read_text(encoding='utf-8').strip() if fix_path and Path(fix_path).exists() else ''
                doc.add_paragraph('Prompt:')
                doc.add_paragraph(prompt)
                if base_text:
                    doc.add_paragraph('Normal:')
                    doc.add_paragraph(base_text)
                if fix_text:
                    doc.add_paragraph('Fixed-Up (Cadence-Controlled):')
                    doc.add_paragraph(fix_text)
        doc.save(str(docx_path))
        print(f'DOCX written to {docx_path}')
    except Exception as e:
        print(f'[WARN] Could not write DOCX ({e}). Try: pip install python-docx')

    # Build single-page HTML
    build_final_html(args.model, images, latest, Path('reports/final/report.html'))


if __name__ == '__main__':
    main()
